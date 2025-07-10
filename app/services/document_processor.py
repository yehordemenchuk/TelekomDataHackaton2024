import asyncio
import hashlib
import os
import shutil
from typing import Dict, Any, Optional, List, Tuple
import logging
from pathlib import Path
import aiofiles
import magic
# Document processing imports
import PyPDF2
from docx import Document as DocxDocument
import docx
import zipfile
import tempfile

from ..config import settings
from ..models.schemas import DocumentType

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing and extracting content from various document formats."""
    
    SUPPORTED_FORMATS = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword',
        'txt': 'text/plain',
        'rtf': 'application/rtf'
    }
    
    def __init__(self):
        self.storage_path = Path(settings.document_storage_path)
        self.storage_path.mkdir(exist_ok=True)
    
    async def process_uploaded_document(
        self,
        file_content: bytes,
        filename: str,
        title: str,
        document_type: DocumentType,
        owner_id: int,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Tuple[Dict[str, Any], str]:
        """Process an uploaded document and extract its content."""
        try:
            # Validate file format
            file_format = await self._detect_file_format(file_content, filename)
            if not file_format:
                raise ValueError(f"Unsupported file format: {filename}")
            
            # Save file to storage
            file_path = await self._save_file(file_content, filename, owner_id)
            
            # Extract text content
            content = await self._extract_content(file_content, file_format)
            
            # Extract metadata from file
            extracted_metadata = await self.extract_document_metadata(file_content, file_format)
            
            # Combine with provided metadata
            final_metadata = metadata or {}
            final_metadata.update(extracted_metadata)
            
            # Create document info
            document_info = {
                "title": title,
                "filename": filename,
                "file_path": str(file_path),
                "document_type": document_type.value,
                "owner_id": owner_id,
                "metadata": final_metadata,
                "file_size": len(file_content)
            }
            
            logger.info(f"Processed document: {filename}")
            return document_info, content
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            raise
    
    async def _detect_file_format(self, file_content: bytes, filename: str) -> Optional[str]:
        """Detect the file format based on content and filename."""
        try:
            # Try to detect MIME type from content
            mime_type = magic.from_buffer(file_content, mime=True)
            
            # Check against supported formats
            for format_name, format_mime in self.SUPPORTED_FORMATS.items():
                if mime_type == format_mime:
                    return format_name
            
            # Fallback to file extension
            file_ext = Path(filename).suffix.lower().lstrip('.')
            if file_ext in self.SUPPORTED_FORMATS:
                return file_ext
            
            return None
            
        except Exception as e:
            logger.error(f"Error detecting file format: {str(e)}")
            return None
    
    async def _save_file(self, file_content: bytes, filename: str, owner_id: int) -> Path:
        """Save file to storage with organized directory structure."""
        try:
            # Create user-specific directory
            user_dir = self.storage_path / f"user_{owner_id}"
            user_dir.mkdir(exist_ok=True)
            
            # Generate unique filename
            timestamp = int(asyncio.get_event_loop().time())
            file_extension = Path(filename).suffix
            unique_filename = f"{timestamp}_{filename}"
            
            file_path = user_dir / unique_filename
            
            # Save file asynchronously
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(file_content)
            
            return file_path
            
        except Exception as e:
            logger.error(f"Error saving file {filename}: {str(e)}")
            raise
    
    async def _extract_content(self, file_content: bytes, file_format: str) -> str:
        """Extract text content from file based on format."""
        try:
            if file_format == 'pdf':
                return await self._extract_pdf_content(file_content)
            elif file_format == 'docx':
                return await self._extract_docx_content(file_content)
            elif file_format == 'txt':
                return file_content.decode('utf-8', errors='ignore')
            elif file_format == 'rtf':
                return await self._extract_rtf_content(file_content)
            else:
                raise ValueError(f"Unsupported format for content extraction: {file_format}")
                
        except Exception as e:
            logger.error(f"Error extracting content from {file_format}: {str(e)}")
            raise
    
    async def _extract_pdf_content(self, file_content: bytes) -> str:
        """Extract text content from PDF."""
        try:
            # Create a temporary file for PDF processing
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                text_content = ""
                with open(temp_file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text_content += page.extract_text() + "\n"
                
                return text_content.strip()
                
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
                
        except Exception as e:
            logger.error(f"Error extracting PDF content: {str(e)}")
            raise
    
    async def _extract_docx_content(self, file_content: bytes) -> str:
        """Extract text content from DOCX."""
        try:
            # Create a temporary file for DOCX processing
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                doc = DocxDocument(temp_file_path)
                text_content = ""
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_content += cell.text + " "
                        text_content += "\n"
                
                return text_content.strip()
                
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
                
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {str(e)}")
            raise
    
    async def _extract_rtf_content(self, file_content: bytes) -> str:
        """Extract text content from RTF (basic implementation)."""
        try:
            # This is a simplified RTF parser
            # For production, consider using a proper RTF library like striprtf
            
            content = file_content.decode('utf-8', errors='ignore')
            
            # Remove RTF control codes (basic cleanup)
            import re
            
            # Remove RTF control words
            content = re.sub(r'\\[a-z]+\d*\s?', '', content)
            # Remove braces
            content = re.sub(r'[{}]', '', content)
            # Clean up whitespace
            content = re.sub(r'\s+', ' ', content)
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"Error extracting RTF content: {str(e)}")
            raise
    
    async def _extract_content_from_file(self, file_path: str) -> str:
        """Extract content from an existing file."""
        try:
            async with aiofiles.open(file_path, 'rb') as f:
                file_content = await f.read()
            
            file_format = await self._detect_file_format(file_content, file_path)
            if not file_format:
                raise ValueError(f"Cannot determine format for file: {file_path}")
            
            return await self._extract_content(file_content, file_format)
            
        except Exception as e:
            logger.error(f"Error extracting content from file {file_path}: {str(e)}")
            raise
    
    async def extract_document_metadata(self, file_content: bytes, file_format: str) -> Dict[str, Any]:
        """Extract metadata from document."""
        metadata = {
            "file_size": len(file_content),
            "format": file_format
        }
        
        try:
            if file_format == 'pdf':
                metadata.update(await self._extract_pdf_metadata(file_content))
            elif file_format == 'docx':
                metadata.update(await self._extract_docx_metadata(file_content))
            
        except Exception as e:
            logger.warning(f"Could not extract metadata: {str(e)}")
        
        return metadata
    
    async def _extract_pdf_metadata(self, file_content: bytes) -> Dict[str, Any]:
        """Extract metadata from PDF."""
        metadata = {}
        
        try:
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                with open(temp_file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    metadata["page_count"] = len(pdf_reader.pages)
                    
                    if pdf_reader.metadata:
                        if '/Title' in pdf_reader.metadata:
                            metadata["title"] = pdf_reader.metadata['/Title']
                        if '/Author' in pdf_reader.metadata:
                            metadata["author"] = pdf_reader.metadata['/Author']
                        if '/Creator' in pdf_reader.metadata:
                            metadata["creator"] = pdf_reader.metadata['/Creator']
                        if '/CreationDate' in pdf_reader.metadata:
                            metadata["creation_date"] = str(pdf_reader.metadata['/CreationDate'])
                
            finally:
                os.unlink(temp_file_path)
                
        except Exception as e:
            logger.warning(f"Error extracting PDF metadata: {str(e)}")
        
        return metadata
    
    async def _extract_docx_metadata(self, file_content: bytes) -> Dict[str, Any]:
        """Extract metadata from DOCX."""
        metadata = {}
        
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                doc = DocxDocument(temp_file_path)
                
                # Count pages (approximate)
                paragraph_count = len(doc.paragraphs)
                metadata["paragraph_count"] = paragraph_count
                metadata["estimated_pages"] = max(1, paragraph_count // 20)  # Rough estimate
                
                # Extract core properties if available
                if hasattr(doc, 'core_properties'):
                    core_props = doc.core_properties
                    if core_props.title:
                        metadata["title"] = core_props.title
                    if core_props.author:
                        metadata["author"] = core_props.author
                    if core_props.created:
                        metadata["creation_date"] = str(core_props.created)
                    if core_props.modified:
                        metadata["modified_date"] = str(core_props.modified)
                
            finally:
                os.unlink(temp_file_path)
                
        except Exception as e:
            logger.warning(f"Error extracting DOCX metadata: {str(e)}")
        
        return metadata
    
    async def delete_document_file(self, file_path: str):
        """Delete a document file from storage."""
        try:
            if os.path.exists(file_path):
                os.unlink(file_path)
                logger.info(f"Deleted file: {file_path}")
            
        except Exception as e:
            logger.error(f"Error deleting file {file_path}: {str(e)}")
            raise
    
    async def validate_document_content(self, content: str, document_type: DocumentType) -> Dict[str, Any]:
        """Validate document content based on type."""
        validation_results = {
            "is_valid": True,
            "issues": [],
            "word_count": len(content.split()),
            "character_count": len(content),
            "estimated_reading_time": len(content.split()) / 200  # Average reading speed
        }
        
        # Basic content validation
        if len(content.strip()) < 10:
            validation_results["is_valid"] = False
            validation_results["issues"].append("Document content is too short")
        
        # Type-specific validation
        if document_type == DocumentType.CONTRACT:
            if "party" not in content.lower() and "parties" not in content.lower():
                validation_results["issues"].append("Contract should mention parties involved")
            
            if "consideration" not in content.lower():
                validation_results["issues"].append("Contract should specify consideration")
        
        elif document_type == DocumentType.WILL:
            if "will" not in content.lower() and "testament" not in content.lower():
                validation_results["issues"].append("Will document should contain 'will' or 'testament'")
            
            if "executor" not in content.lower():
                validation_results["issues"].append("Will should name an executor")
        
        elif document_type == DocumentType.NDA:
            if "confidential" not in content.lower():
                validation_results["issues"].append("NDA should mention confidentiality")
            
            if "disclosure" not in content.lower():
                validation_results["issues"].append("NDA should address disclosure terms")
        
        if validation_results["issues"]:
            validation_results["is_valid"] = False
        
        return validation_results
    
    async def get_storage_statistics(self, owner_id: Optional[int] = None) -> Dict[str, Any]:
        """Get file storage statistics."""
        try:
            total_size = 0
            file_count = 0
            
            if owner_id:
                user_dir = self.storage_path / f"user_{owner_id}"
                if user_dir.exists():
                    for file_path in user_dir.rglob("*"):
                        if file_path.is_file():
                            total_size += file_path.stat().st_size
                            file_count += 1
            else:
                # Calculate for all users
                for file_path in self.storage_path.rglob("*"):
                    if file_path.is_file():
                        total_size += file_path.stat().st_size
                        file_count += 1
            
            return {
                "total_files": file_count,
                "total_storage_size": total_size,
                "storage_size_mb": round(total_size / (1024 * 1024), 2),
                "storage_path": str(self.storage_path)
            }
            
        except Exception as e:
            logger.error(f"Error getting storage statistics: {str(e)}")
            raise 